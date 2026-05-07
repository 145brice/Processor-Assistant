"""Generate filled HOI Request and Title Request docs from a loan record."""
import json
import os
import re
from datetime import datetime

TEMPLATES_DIR = r"c:\Users\user\OneDrive\Desktop\Templates"
OUTPUT_ROOT = r"c:\Users\user\OneDrive\Desktop\processor-traien\Processor-Assistant\generated_docs"
_CLAUSES_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "mortgagee_clauses.json")


def load_lender_clauses() -> dict:
    """Returns {lender_key: clause_text}. Empty dict if file missing/malformed."""
    try:
        with open(_CLAUSES_FILE, "r", encoding="utf-8") as f:
            return (json.load(f) or {}).get("lenders", {}) or {}
    except Exception:
        return {}


def get_lender_names() -> list[str]:
    return sorted(load_lender_clauses().keys())


def get_mortgagee_clause(lender: str, loan_number: str = "") -> str:
    """Return the clause text for the given lender, with {LOAN_NUMBER} substituted.
    Returns the placeholder text if lender is unknown/empty."""
    if not lender:
        return "[MORTGAGEE CLAUSE — fill in per lender]"
    clauses = load_lender_clauses()
    text = clauses.get(lender, "")
    if not text:
        return "[MORTGAGEE CLAUSE — fill in per lender]"
    return text.replace("{LOAN_NUMBER}", str(loan_number or ""))

TEMPLATE_FIELDS = {
    "HOI Request.docx": {
        "Borrower(s):": "borrower_name",
        "Property:": "property_address",
        "Estimated Funding:": "funding_date",
        "Loan Number:": "loan_num",
        "Purchase Price/Appraised Value:": "purchase_price",
        "Est. Total Loan Amount:": "loan_amount",
        "Type:": "loan_type",
        "Loan Officer:": "loan_officer",
        "Loan Processor:": "loan_processor",
    },
    "Title Request copy.docx": {
        "Borrower(s):": "borrower_name",
        "Property:": "property_address",
        "Estimated Funding:": "funding_date",
        "Loan Number:": "loan_num",
        "Purchase/Appraised Price:": "purchase_price",
        "Est. Total Loan Amount:": "loan_amount",
        "Type:": "loan_type",
        "Loan Officer:": "loan_officer",
        "Loan Processor:": "loan_processor",
    },
}

MORTGAGEE_PARAGRAPH_INDEX = {
    "HOI Request.docx": 15,
    "Title Request copy.docx": 12,
}


def _from_docs(loan: dict) -> dict:
    """Pull fallback values from attached docs' extracted_data."""
    out: dict = {}
    for doc in loan.get("documents", []) or []:
        data = doc.get("extracted_data") or {}
        dtype = doc.get("doc_type", "")
        if dtype == "Purchase Contract":
            prop = (data.get("property") or {}).get("address")
            price = (data.get("transaction") or {}).get("purchase_price")
            if prop and not out.get("property_address"):
                out["property_address"] = prop
            if price and not out.get("purchase_price"):
                out["purchase_price"] = price
        elif dtype == "1003 Application":
            loan_info = data.get("loan") or {}
            prop = data.get("property_address") or loan_info.get("property_address")
            if prop and not out.get("property_address"):
                out["property_address"] = prop
            if loan_info.get("amount") and not out.get("loan_amount"):
                out["loan_amount"] = loan_info["amount"]
            if loan_info.get("purpose") and not out.get("loan_type"):
                out["loan_type"] = loan_info["purpose"]
            lo = data.get("loan_officer")
            if lo and not out.get("loan_officer"):
                out["loan_officer"] = lo.get("name") if isinstance(lo, dict) else lo
    return out


def build_context(loan: dict) -> dict:
    contacts = loan.get("contacts", {}) or {}
    borrower = contacts.get("borrower", {}) or {}
    buyer = contacts.get("buyer", {}) or {}
    name = borrower.get("name") or buyer.get("name") or loan.get("borrower", "")

    from_docs = _from_docs(loan)

    def pick(key: str, placeholder: str) -> str:
        return loan.get(key) or from_docs.get(key) or placeholder

    lender = loan.get("lender", "") or ""
    return {
        "borrower_name": name or "[BORROWER NAME]",
        "property_address": pick("property_address", "[PROPERTY ADDRESS]"),
        "funding_date": loan.get("closing_date", "") or "[FUNDING DATE]",
        "loan_num": loan.get("loan_num", "") or "[LOAN NUMBER]",
        "purchase_price": pick("purchase_price", "[PURCHASE PRICE]"),
        "loan_amount": pick("loan_amount", "[LOAN AMOUNT]"),
        "loan_type": pick("loan_type", "[PURCHASE/REFI]"),
        "loan_officer": pick("loan_officer", "[LOAN OFFICER]"),
        "loan_processor": loan.get("loan_processor") or "Brice Leasure",
        "lender": lender,
        "mortgagee_clause": get_mortgagee_clause(lender, loan.get("loan_num", "")),
    }


def _replace_after_label(paragraph, label: str, value: str):
    """Replace whatever comes after 'Label:' with the new value, keeping the label."""
    text = paragraph.text
    if label not in text:
        return False
    new_text = re.sub(
        rf"({re.escape(label)})\s*.*$",
        rf"\1 {value}",
        text,
    )
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)
    return True


def fill_template(template_name: str, context: dict, output_path: str):
    from docx import Document

    src = os.path.join(TEMPLATES_DIR, template_name)
    doc = Document(src)
    field_map = TEMPLATE_FIELDS[template_name]

    for paragraph in doc.paragraphs:
        for label, ctx_key in field_map.items():
            if label in paragraph.text:
                _replace_after_label(paragraph, label, str(context[ctx_key]))
                break

    mortgagee_idx = MORTGAGEE_PARAGRAPH_INDEX[template_name]
    clause_text = context.get("mortgagee_clause") or "[MORTGAGEE CLAUSE — fill in per lender]"
    if mortgagee_idx < len(doc.paragraphs):
        p = doc.paragraphs[mortgagee_idx]
        for run in p.runs:
            run.text = ""
        # Multi-line clauses: use line breaks (docx requires \n inside a run via add_break)
        lines = clause_text.split("\n")
        if p.runs:
            p.runs[0].text = lines[0]
            for extra in lines[1:]:
                p.runs[0].add_break()
                p.add_run(extra)
        else:
            r = p.add_run(lines[0])
            for extra in lines[1:]:
                r.add_break()
                p.add_run(extra)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_for_loan(loan: dict, templates=None) -> list[str]:
    templates = templates or list(TEMPLATE_FIELDS.keys())
    ctx = build_context(loan)
    loan_id = loan.get("id", "unknown")
    out_dir = os.path.join(OUTPUT_ROOT, str(loan_id))
    written = []
    for tpl in templates:
        safe_borrower = re.sub(r"[^A-Za-z0-9_-]+", "_", ctx["borrower_name"])[:40]
        out_name = tpl.replace(".docx", f"_{safe_borrower}.docx")
        written.append(fill_template(tpl, ctx, os.path.join(out_dir, out_name)))
    return written


if __name__ == "__main__":
    with open("pipeline.json", "r", encoding="utf-8") as f:
        pipeline = json.load(f)
    loan9 = next(l for l in pipeline if l.get("id") == 9)
    paths = generate_for_loan(loan9)
    for p in paths:
        print(f"wrote: {p}")
